from fastapi import FastAPI, UploadFile, File
from fastapi.responses import HTMLResponse, FileResponse
import zipfile
import json
import os
import shutil
from docx import Document

app = FastAPI()

UPLOAD_DIR = "temp_upload"
EXTRACT_DIR = "temp_extract"
OUTPUT_FILE = "output.docx"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(EXTRACT_DIR, exist_ok=True)


# -----------------------------
# Helper functions
# -----------------------------

def get_label(obj):
    if isinstance(obj.get("label"), dict):
        return obj["label"].get("en")

    return (
        obj.get("label")
        or obj.get("title")
        or obj.get("name")
        or obj.get("identifier")
        or "Unnamed"
    )


def get_type(field):
    return (
        field.get("data_type")
        or field.get("type")
        or field.get("field_type")
        or "unknown"
    )


def is_user_field(field):
    name = str(field.get("identifier", "")).lower()

    return not any(x in name for x in [
        "created", "updated", "status", "reference", "id", "internal"
    ])


def get_required(field):
    return field.get("required") or field.get("is_required") or False


def get_help_text(field):
    return field.get("help_text") or field.get("description") or ""


def get_options(field):
    return field.get("options") or field.get("choices") or []


# -----------------------------
# Routes
# -----------------------------

@app.get("/", response_class=HTMLResponse)
def home():
    return """
    <h2>Jadu Export → Form Spec</h2>
    <form action="/upload" method="post" enctype="multipart/form-data">
      <input type="file" name="file" accept=".zip" required />
      <button type="submit">Generate spec</button>
    </form>
    """


@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    try:
        # Save file
        zip_path = os.path.join(UPLOAD_DIR, file.filename)
        with open(zip_path, "wb") as f:
            f.write(await file.read())

        # Reset extract folder
        if os.path.exists(EXTRACT_DIR):
            shutil.rmtree(EXTRACT_DIR)
        os.makedirs(EXTRACT_DIR, exist_ok=True)

        # Extract ZIP
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(EXTRACT_DIR)

        # Find valid JSON
        json_file = None

        for root, _, files in os.walk(EXTRACT_DIR):
            for fname in files:
                if fname.endswith(".json"):
                    path = os.path.join(root, fname)

                    if os.path.getsize(path) == 0:
                        continue

                    try:
                        with open(path, encoding="utf-8", errors="replace") as f:
                            data = json.load(f)

                        if "resources" in data:
                            json_file = path
                            break

                    except Exception:
                        continue

            if json_file:
                break

        if not json_file:
            return {"error": "No valid Jadu JSON file found"}

        # Load JSON
        with open(json_file, encoding="utf-8", errors="replace") as f:
            data = json.load(f)

        resources = data.get("resources", {})

        # Group resources
        grouped = {}
        for key, value in resources.items():
            rtype = key.split("::")[0]
            grouped.setdefault(rtype, []).append(value)

        # -----------------------------
        # Build document
        # -----------------------------

        doc = Document()
        doc.add_heading("Form specification", 0)

        # Overview
        doc.add_heading("Overview", 1)
        doc.add_paragraph("Generated from Jadu export. This describes the user-facing form and workflow.")

        # -----------------------------
        # Fields
        # -----------------------------
        doc.add_heading("Form fields", 1)

        for field in grouped.get("case-field", []):
            if not isinstance(field, dict):
                continue

            if not is_user_field(field):
                continue

            label = get_label(field)
            ftype = get_type(field)
            required = get_required(field)
            help_text = get_help_text(field)
            options = get_options(field)

            doc.add_heading(str(label), 2)
            doc.add_paragraph(f"Type: {ftype}")
            doc.add_paragraph(f"Required: {'Yes' if required else 'No'}")

            if help_text:
                doc.add_paragraph(f"Help text: {help_text}")

            if options and isinstance(options, list):
                doc.add_paragraph("Options:")
                for opt in options:
                    doc.add_paragraph(f"- {opt}", style='List Bullet')

        # -----------------------------
        # Workflow (statuses)
        # -----------------------------
        doc.add_heading("Case workflow", 1)

        for status in grouped.get("case-status", []):
            if not isinstance(status, dict):
                continue

            name = get_label(status)
            doc.add_paragraph(str(name))

        # -----------------------------
        # Transitions
        # -----------------------------
        doc.add_heading("Transitions", 1)

        for transition in grouped.get("case-transition", []):
            if not isinstance(transition, dict):
                continue

            name = get_label(transition)

            from_status = transition.get("from_status")
            to_status = transition.get("to_status")

            doc.add_paragraph(str(name))

            if from_status or to_status:
                doc.add_paragraph(
                    f"From: {from_status or 'Unknown'} → To: {to_status or 'Unknown'}"
                )

        # -----------------------------
        # Rules
        # -----------------------------
        doc.add_heading("Rules", 1)

        for rule in grouped.get("case-rule", []):
            if not isinstance(rule, dict):
                continue

            name = get_label(rule)
            doc.add_paragraph(str(name))

        # -----------------------------
        # Emails
        # -----------------------------
        doc.add_heading("Notifications", 1)

        for email in grouped.get("alert-email-template", []):
            if not isinstance(email, dict):
                continue

            subject = email.get("subject") or "No subject"
            body = email.get("body") or ""

            doc.add_heading(subject, 2)

            if body:
                doc.add_paragraph(body[:300])

        # Save
        doc.save(OUTPUT_FILE)

        # Cleanup
        os.remove(zip_path)

        return FileResponse(OUTPUT_FILE, filename="form-spec.docx")

    except Exception as e:
        return {"error": str(e)}
